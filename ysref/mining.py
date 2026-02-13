import glob
import json
import os
import re
import fitz
fitz.TOOLS.mupdf_display_errors(False)
import docx
import openpyxl
from bs4 import BeautifulSoup

import ysref.util

# 
def updateMinedWord(mined, query, txt, file):
  matched = re.findall(query, txt)
  for wrd in matched:
    pos = txt.find(wrd)
    sz = len(wrd)
    peripheral = f"... {txt[(pos-6 if 6 < pos else 0) : (pos + sz + 6 if ((pos + sz + 6) < len(txt)) else len(txt))]} ..."
    if wrd in mined:
      if file in mined[wrd]:
        mined[wrd][file].append(peripheral)
      else:
        mined[wrd][file] = [peripheral]
    else:
      mined[wrd] = {file: [peripheral]}

# Mining from text file
def mineFromTXT(mined, query, path):
  (dir,file) = os.path.split(path)
  txt = ''
  with open(path) as f:
    txt = f.read()
  updateMinedWord(mined, query, txt, file)
  
# Define PDF parser func.
def parsePDF(path, dir):
  # Make dir.
  os.makedirs(dir, exist_ok=True)
  # Open PDF
  doc = fitz.open(path)
  # Prepare container
  pages = []
  # Extract elements from each page
  for p in range(len(doc)):
    # Prepare container
    data = {
      'page': p+1,
      'text': '',
      'img': []
    }
    # Load page
    page = doc[p]
    # Extract texts
    text = page.get_text("text")
    data['text'] = text
    # Extract images
    for i, img in enumerate(page.get_images(full=True)):
      # Get image ID
      imgid = img[0]
      # Image Obj.
      image = doc.extract_image(imgid)
      # Bynary data
      imgdata = image['image']
      # Save image
      imgfile = f"{dir}/{p + 1}_{i + 1}.{image['ext']}"
      data['img'].append(imgfile)
      with open(imgfile, "wb") as f:
        f.write(imgdata)
      # Store
      pages.append(data)
  # Close
  doc.close()
  # Return
  return pages

# Mining from PDF file
def mineFromPDF(mined, query, path):
  """
   Mining from a PDF document 

    Args:
        mined (dict): Container to save result.
        query (str): The regex pattern to search for.
        path (str): Input file path.
  """
  # Path separation
  (dir,file) = os.path.split(path)
  (name,ext) = os.path.splitext(file)
  # Call parser
  contents = parsePDF(path, os.path.join(dir, 'contents'))
  # Save parsed data
  json.dump(contents, open(f"{name}_extract.json", "w"))
  # Check each page
  for page in contents:
    updateMinedWord(mined, query, page['text'], file)

# Mining from docx file
def mineFromDOC(mined, query, path):
  """
   Mining from a document 

    Args:
        mined (dict): Container to save result.
        query (str): The regex pattern to search for.
        path (str): Input file path.
  """
  # Open docment
  (dir,file) = os.path.split(path)
  doc = docx.Document(path)
  # Check each paragraph
  for para in doc.paragraphs:
    updateMinedWord(mined, query, para.text, file)
  # Check embeded table
  for table_index, table in enumerate(doc.tables):
    for row_index, row in enumerate(table.rows):
      for cell_index, cell in enumerate(row.cells):
        updateMinedWord(mined, query, cell.text, file)
  
# Mining from xlsx file
def mineFromXLS(mined, query, path, max_fsize=(3<<20)):
  """
   Mining from a table data

    Args:
        mined (dict): Container to save result.
        query (str): The regex pattern to search for.
        path (str): Input file path.
  """
  # Ignore huge dataset file
  if max_fsize < os.path.getsize(path):
    return
  # Open workbook
  (dir,file) = os.path.split(path)
  workbook = openpyxl.load_workbook(path)
  # Check each cell
  for ws in workbook.worksheets:
    for row in ws.iter_rows():
      for cell in row:
        if cell.value:
          updateMinedWord(mined, query, str(cell.value), file)

# Mining from file
def mineWord(result, query, path, verbose=False):
  """
   Mining from a single file

    Args:
        result (dict): Container to save result.
        query (str): The regex pattern to search for.
        path (str): Input file path.
  """
  if path.endswith('htm') or path.endswith('html') or path.endswith('xml') or path.endswith('txt'):
    mineFromTXT(result, query, path)
  if path.endswith('pdf') or path.endswith('PDF'):
    mineFromPDF(result, query, path)
  elif path.endswith('docx'):
    mineFromDOC(result, query, path)
  elif path.endswith('xlsx'):
    mineFromXLS(result, query, path)

# Mining from directory
def mineWordFrom(result, query, dir, verbose=False):
  """
    Mining regex query from files in the specified directory recursively.

    Args:
        result (dict): Container to save result.
        query (str): The regex pattern to search for.
        dir (str): Input directory.

  """
  ## List all files in the input directory
  os.chdir(dir)
  files = glob.glob(os.path.join(dir, "*"))
  ## Expand all the compressed files
  for f in files:
    if os.path.isfile(f):
      if f.endswith('.tar.gz'):
        res = ysref.util.execCmd(f"tar -xvzf '{f}'")
        if res[0]:
          ysref.util.execCmd(f"rm '{f}'")
        else:
          print(res[1], '\n', res[2])
      elif f.endswith('.gz'):
        res = ysref.util.execCmd(f"gunzip '{f}'")
        if res[0]:
          ysref.util.execCmd(f"rm '{f}'")
        else:
          print(res[1], '\n', res[2])
      elif f.endswith('.zip'):
        res = ysref.util.execCmd(f"unzip '{f}'")
        if res[0]:
          ysref.util.execCmd(f"rm '{f}'")
        else:
          print(res[1], '\n', res[2])

  ## Re-listing all files in the input directory
  files = glob.glob(os.path.join(dir, "*"))
  for f in files:
    try:
      ## Mining from a NOT empty file
      if os.path.isfile(f) and 0 < os.path.getsize(f):
        mineWord(result, query, f, verbose=verbose)
      ## Mining from a directory (recursive)
      elif os.path.isdir(f):
        mineWordFrom(result, query, f, verbose=verbose)
    except Exception as e:
      print(e)
    